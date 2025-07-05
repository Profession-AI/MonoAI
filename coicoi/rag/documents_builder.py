import os
import uuid
from typing import List, Dict, Tuple, Optional
from coicoi.tools.webscraping import WebScraping

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class DocumentsBuilder:
    """
    A utility class for building document collections from various sources.
    
    This class provides methods to extract text content from files and web pages,
    split the content into manageable chunks with configurable size and overlap,
    and prepare the data for storage in vector databases.
    
    The DocumentsBuilder is designed to work seamlessly with the RAG system,
    producing output that can be directly used with vector database operations.
    
    Features:
    - File-based document extraction with UTF-8 encoding support
    - Web scraping with multiple engine options (requests, tavily, selenium)
    - Word document extraction (.doc and .docx formats)
    - PDF document extraction with metadata
    - Intelligent text chunking with word boundary preservation
    - Configurable chunk size and overlap parameters
    - Rich metadata generation for each document chunk
    - Unique ID generation for database storage
    
    Attributes:
        _chunk_size (int): Maximum size of each text chunk in characters
        _chunk_overlap (int): Number of characters to overlap between chunks
    
    Examples:
    --------
    Basic usage with file processing:
    
    ```python
    # Initialize with default chunk settings
    builder = DocumentsBuilder()
    
    # Process a text file
    documents, metadatas, ids = builder.from_file("document.txt")
    
    # Add to vector database
    vector_db.add(documents, metadatas, ids)
    ```
    
    Custom chunk settings:
    
    ```python
    # Initialize with custom chunk parameters
    builder = DocumentsBuilder(chunk_size=500, chunk_overlap=50)
    
    # Process multiple files
    for file_path in ["doc1.txt", "doc2.txt"]:
        documents, metadatas, ids = builder.from_file(file_path)
        vector_db.add(documents, metadatas, ids)
    ```
    
    Web scraping with different engines:
    
    ```python
    # Basic web scraping
    documents, metadatas, ids = builder.from_url("https://example.com")
    
    # Advanced scraping with Tavily
    documents, metadatas, ids = builder.from_url(
        "https://example.com",
        engine="tavily",
        deep=True
    )
    
    # JavaScript-heavy sites with Selenium
    documents, metadatas, ids = builder.from_url(
        "https://spa-example.com",
        engine="selenium"
    )
    ```
    
    Word document processing:
    
    ```python
    # Process Word documents
    documents, metadatas, ids = builder.from_doc("document.docx")
    documents, metadatas, ids = builder.from_doc("document.doc")
    
    # Process with custom extraction method
    documents, metadatas, ids = builder.from_doc(
        "document.docx",
        extraction_method="docx2txt"
    )
    ```
    
    PDF document processing:
    
    ```python
    # Process PDF documents
    documents, metadatas, ids = builder.from_pdf("document.pdf")
    
    # Process with page range
    documents, metadatas, ids = builder.from_pdf(
        "document.pdf",
        page_range=(1, 10)  # Extract pages 1-10
    )
    ```
    
    Notes:
    ------
    - Chunk overlap helps maintain context between chunks
    - Word boundaries are preserved when possible to avoid cutting words
    - Empty chunks are automatically filtered out
    - All text is processed as UTF-8
    - Word document processing requires python-docx and python-docx2txt packages
    - PDF processing requires PyPDF2 package
    """

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 100):
        """
        Initialize the DocumentsBuilder with chunking parameters.
        
        Parameters:
        -----------
        chunk_size : int, default=1000
            Maximum size of each text chunk in characters. Larger chunks
            preserve more context but may be less focused for retrieval.
            
        chunk_overlap : int, default=100
            Number of characters to overlap between consecutive chunks.
            Overlap helps maintain context and improves retrieval quality
            for queries that span chunk boundaries.
            
        Examples:
        ---------
        ```python
        # Default settings (good for most use cases)
        builder = DocumentsBuilder()
        
        # Small chunks for precise retrieval
        builder = DocumentsBuilder(chunk_size=500, chunk_overlap=50)
        
        # Large chunks for context preservation
        builder = DocumentsBuilder(chunk_size=2000, chunk_overlap=200)
        ```
        
        Notes:
        ------
        - chunk_overlap should typically be 10-20% of chunk_size
        - Very small chunks (< 200 chars) may lose context
        - Very large chunks (> 3000 chars) may be less focused for retrieval
        """
        self._chunk_size = chunk_size
        self._chunk_overlap = chunk_overlap

    def from_file(self, file_path: str) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Read a file and split it into chunks with specified size and overlap.
        
        This method reads a text file from the filesystem, splits its content
        into chunks according to the configured parameters, and generates
        metadata and unique IDs for each chunk.
        
        Parameters:
        -----------
        file_path : str
            Path to the text file to read. The file must exist and be
            readable. UTF-8 encoding is assumed.
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The text content split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              file information and chunk details
            - List of unique IDs: UUID strings for each chunk
            
        Raises:
        -------
        FileNotFoundError
            If the specified file does not exist or is not accessible.
            
        UnicodeDecodeError
            If the file cannot be decoded as UTF-8.
            
        Examples:
        ---------
        ```python
        # Process a single file
        documents, metadatas, ids = builder.from_file("article.txt")
        
        # Access metadata information
        for i, metadata in enumerate(metadatas):
            print(f"Chunk {i+1}:")
            print(f"  File: {metadata['file_name']}")
            print(f"  Size: {metadata['chunk_size']} characters")
            print(f"  Position: {metadata['chunk_index'] + 1}/{metadata['total_chunks']}")
        ```
        
        Notes:
        ------
        - File is read entirely into memory before processing
        - Empty files will return empty lists
        - File path is stored in metadata for traceability
        - Chunk indexing starts at 0
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Read the file content
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_doc(self, file_path: str, extraction_method: str = "auto") -> Tuple[List[str], List[Dict], List[str]]:
        """
        Extract text from Word documents (.doc and .docx files) and split into chunks.
        
        This method supports both .doc and .docx formats using different extraction
        methods. For .docx files, it can use either python-docx or docx2txt libraries.
        For .doc files, it uses docx2txt which can handle the older format.
        
        Parameters:
        -----------
        file_path : str
            Path to the Word document (.doc or .docx file). The file must exist
            and be readable.
            
        extraction_method : str, default="auto"
            The method to use for text extraction:
            - "auto": Automatically choose the best method based on file extension
            - "docx": Use python-docx library (only for .docx files)
            - "docx2txt": Use docx2txt library (works for both .doc and .docx)
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The extracted text split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              file information, document properties, and chunk details
            - List of unique IDs: UUID strings for each chunk
            
        Raises:
        -------
        FileNotFoundError
            If the specified file does not exist or is not accessible.
            
        ValueError
            If the file is not a supported Word document format or if the
            required extraction method is not available.
            
        ImportError
            If the required libraries for the chosen extraction method are not installed.
            
        Examples:
        ---------
        ```python
        # Process a .docx file with automatic method selection
        documents, metadatas, ids = builder.from_doc("document.docx")
        
        # Process a .doc file
        documents, metadatas, ids = builder.from_doc("document.doc")
        
        # Force specific extraction method
        documents, metadatas, ids = builder.from_doc(
            "document.docx",
            extraction_method="docx2txt"
        )
        
        # Access document metadata
        for metadata in metadatas:
            print(f"File: {metadata['file_name']}")
            print(f"Format: {metadata['document_format']}")
            print(f"Extraction method: {metadata['extraction_method']}")
        ```
        
        Notes:
        ------
        - .docx files are the modern Word format (Office 2007+)
        - .doc files are the legacy Word format (Office 97-2003)
        - python-docx provides better structure preservation for .docx files
        - docx2txt works with both formats but may lose some formatting
        - Document properties (title, author, etc.) are extracted when available
        - Images and complex formatting are not preserved in the extracted text
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file extension and validate
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension not in ['.doc', '.docx']:
            raise ValueError(f"Unsupported file format: {file_extension}. Only .doc and .docx files are supported.")
        
        # Determine extraction method
        if extraction_method == "auto":
            if file_extension == '.docx' and DOCX_AVAILABLE:
                extraction_method = "docx"
            else:
                extraction_method = "docx2txt"
        
        # Extract text based on method
        if extraction_method == "docx":
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx library is required for 'docx' extraction method. Install with: pip install python-docx")
            if file_extension != '.docx':
                raise ValueError("'docx' extraction method only supports .docx files")
            text, doc_properties = self._extract_with_docx(file_path)
        elif extraction_method == "docx2txt":
            if not DOCX2TXT_AVAILABLE:
                raise ImportError("docx2txt library is required for 'docx2txt' extraction method. Install with: pip install python-docx2txt")
            text, doc_properties = self._extract_with_docx2txt(file_path)
        else:
            raise ValueError(f"Unsupported extraction method: {extraction_method}")
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'document_format': file_extension[1:],  # Remove the dot
                'extraction_method': extraction_method,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            # Add document properties if available
            if doc_properties:
                metadata.update(doc_properties)
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_pdf(self, file_path: str, page_range: Optional[Tuple[int, int]] = None) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Extract text from PDF documents and split into chunks.
        
        This method extracts text content from PDF files using PyPDF2 library.
        It supports extracting all pages or a specific range of pages, and
        preserves page information in the metadata.
        
        Parameters:
        -----------
        file_path : str
            Path to the PDF file. The file must exist and be readable.
            
        page_range : Tuple[int, int], optional
            Range of pages to extract (start_page, end_page), where pages are
            1-indexed. If None, all pages are extracted.
            Example: (1, 5) extracts pages 1 through 5.
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The extracted text split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              file information, PDF properties, page information, and chunk details
            - List of unique IDs: UUID strings for each chunk
            
        Raises:
        -------
        FileNotFoundError
            If the specified file does not exist or is not accessible.
            
        ValueError
            If the file is not a valid PDF or if the page range is invalid.
            
        ImportError
            If PyPDF2 library is not installed.
            
        Examples:
        ---------
        ```python
        # Process entire PDF
        documents, metadatas, ids = builder.from_pdf("document.pdf")
        
        # Process specific page range
        documents, metadatas, ids = builder.from_pdf(
            "document.pdf",
            page_range=(1, 10)  # Pages 1-10
        )
        
        # Process single page
        documents, metadatas, ids = builder.from_pdf(
            "document.pdf",
            page_range=(5, 5)  # Only page 5
        )
        
        # Access PDF metadata
        for metadata in metadatas:
            print(f"File: {metadata['file_name']}")
            print(f"Page: {metadata.get('page_number', 'N/A')}")
            print(f"Total pages: {metadata.get('total_pages', 'N/A')}")
            print(f"PDF title: {metadata.get('pdf_title', 'N/A')}")
        ```
        
        Notes:
        ------
        - Page numbers are 1-indexed (first page is page 1)
        - Text extraction quality depends on the PDF structure
        - Scanned PDFs may not extract text properly
        - PDF metadata (title, author, etc.) is extracted when available
        - Page information is preserved in chunk metadata
        - Images and complex formatting are not preserved
        """
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 library is required for PDF processing. Install with: pip install PyPDF2")
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Validate file extension
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension != '.pdf':
            raise ValueError(f"Unsupported file format: {file_extension}. Only .pdf files are supported.")
        
        # Extract text and metadata from PDF
        text, pdf_properties, page_info = self._extract_from_pdf(file_path, page_range)
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'document_format': 'pdf',
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            # Add PDF properties if available
            if pdf_properties:
                metadata.update(pdf_properties)
            
            # Add page information if available
            if page_info:
                metadata.update(page_info)
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids

    def from_url(self, url: str, engine: str = "requests", deep: bool = False) -> Tuple[List[str], List[Dict], List[str]]:
        """
        Scrape content from a URL and split it into chunks with specified size and overlap.
        
        This method uses web scraping to extract text content from a webpage,
        then processes the content using the same chunking logic as file processing.
        Multiple scraping engines are supported for different types of websites.
        
        Parameters:
        -----------
        url : str
            The URL to scrape. Must be a valid HTTP/HTTPS URL.
            
        engine : str, default="requests"
            The web scraping engine to use:
            - "requests": Simple HTTP requests (fast, good for static content)
            - "tavily": Advanced web scraping with better content extraction
            - "selenium": Full browser automation (good for JavaScript-heavy sites)
            
        deep : bool, default=False
            If using the "tavily" engine, whether to use advanced extraction mode.
            Deep extraction provides better content quality but is slower.
            
        Returns:
        --------
        Tuple[List[str], List[Dict], List[str]]
            A tuple containing:
            - List of document chunks (strings): The scraped text split into chunks
            - List of metadata dictionaries: Metadata for each chunk including
              URL information and scraping details
            - List of unique IDs: UUID strings for each chunk
            
        Raises:
        -------
        ValueError
            If the scraping fails or no text content is extracted.
            
        Examples:
        ---------
        ```python
        # Basic web scraping
        documents, metadatas, ids = builder.from_url("https://example.com")
        
        # Advanced scraping with Tavily
        documents, metadatas, ids = builder.from_url(
            "https://blog.example.com",
            engine="tavily",
            deep=True
        )
        
        # JavaScript-heavy site with Selenium
        documents, metadatas, ids = builder.from_url(
            "https://spa.example.com",
            engine="selenium"
        )
        
        # Access scraping metadata
        for metadata in metadatas:
            print(f"Source: {metadata['url']}")
            print(f"Engine: {metadata['scraping_engine']}")
            print(f"Deep extraction: {metadata['deep_extraction']}")
        ```
        
        Notes:
        ------
        - Scraping may take time depending on the engine and website complexity
        - Some websites may block automated scraping
        - Selenium requires Chrome/Chromium to be installed
        - Tavily requires an API key to be configured
        """
        # Initialize WebScraping with specified engine
        scraper = WebScraping(engine=engine, deep=deep)
        
        # Scrape the URL
        result = scraper.scrape(url)
        
        if not result or not result.get("text"):
            raise ValueError(f"Failed to extract text content from URL: {url}")
        
        text = result["text"]
        
        # Split text into chunks
        chunks = self._split_text(text)
        
        # Generate metadata and IDs for each chunk
        documents = []
        metadatas = []
        ids = []
        
        for i, chunk in enumerate(chunks):
            # Generate unique ID
            chunk_id = str(uuid.uuid4())
            
            # Create metadata
            metadata = {
                'url': url,
                'source_type': 'web_page',
                'scraping_engine': engine,
                'deep_extraction': deep,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk)
            }
            
            documents.append(chunk)
            metadatas.append(metadata)
            ids.append(chunk_id)
        
        return documents, metadatas, ids
    
    def _extract_with_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from a .docx file using python-docx library.
        
        Parameters:
        -----------
        file_path : str
            Path to the .docx file
            
        Returns:
        --------
        Tuple[str, Dict]
            A tuple containing the extracted text and document properties
        """
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n\n".join(text_parts)
        
        # Extract document properties
        properties = {}
        core_props = doc.core_properties
        if core_props.title:
            properties['document_title'] = core_props.title
        if core_props.author:
            properties['document_author'] = core_props.author
        if core_props.subject:
            properties['document_subject'] = core_props.subject
        if core_props.created:
            properties['document_created'] = str(core_props.created)
        if core_props.modified:
            properties['document_modified'] = str(core_props.modified)
        
        return text, properties
    
    def _extract_with_docx2txt(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from a Word document (.doc or .docx) using docx2txt library.
        
        Parameters:
        -----------
        file_path : str
            Path to the Word document
            
        Returns:
        --------
        Tuple[str, Dict]
            A tuple containing the extracted text and document properties
        """
        text = docx2txt.process(file_path)
        
        # docx2txt doesn't provide document properties, so return empty dict
        properties = {}
        
        return text, properties
    
    def _extract_from_pdf(self, file_path: str, page_range: Optional[Tuple[int, int]] = None) -> Tuple[str, Dict, Dict]:
        """
        Extract text and metadata from a PDF file using PyPDF2.
        
        Parameters:
        -----------
        file_path : str
            Path to the PDF file
            
        page_range : Tuple[int, int], optional
            Range of pages to extract (start_page, end_page), 1-indexed
            
        Returns:
        --------
        Tuple[str, Dict, Dict]
            A tuple containing the extracted text, PDF properties, and page information
        """
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Get total number of pages
            total_pages = len(pdf_reader.pages)
            
            # Determine page range
            if page_range is None:
                start_page = 1
                end_page = total_pages
            else:
                start_page, end_page = page_range
                # Validate page range
                if start_page < 1 or end_page > total_pages or start_page > end_page:
                    raise ValueError(f"Invalid page range: {page_range}. Pages must be between 1 and {total_pages}")
            
            # Extract text from specified pages
            text_parts = []
            for page_num in range(start_page - 1, end_page):  # Convert to 0-indexed
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            
            # Extract PDF properties
            properties = {}
            if pdf_reader.metadata:
                metadata = pdf_reader.metadata
                if '/Title' in metadata:
                    properties['pdf_title'] = metadata['/Title']
                if '/Author' in metadata:
                    properties['pdf_author'] = metadata['/Author']
                if '/Subject' in metadata:
                    properties['pdf_subject'] = metadata['/Subject']
                if '/Creator' in metadata:
                    properties['pdf_creator'] = metadata['/Creator']
                if '/Producer' in metadata:
                    properties['pdf_producer'] = metadata['/Producer']
                if '/CreationDate' in metadata:
                    properties['pdf_creation_date'] = str(metadata['/CreationDate'])
                if '/ModDate' in metadata:
                    properties['pdf_modification_date'] = str(metadata['/ModDate'])
            
            # Add page information
            page_info = {
                'total_pages': total_pages,
                'extracted_pages_start': start_page,
                'extracted_pages_end': end_page,
                'extracted_pages_count': end_page - start_page + 1
            }
            
            return text, properties, page_info
    
    def _split_text(self, text: str) -> List[str]:
        """
        Split text into chunks with specified size and overlap.
        
        This private method implements the core text chunking algorithm.
        It attempts to break text at word boundaries to avoid cutting words
        in half, while respecting the configured chunk size and overlap parameters.
        
        Parameters:
        -----------
        text : str
            The text content to split into chunks.
            
        Returns:
        --------
        List[str]
            List of text chunks, each with maximum size equal to chunk_size.
            Empty chunks are automatically filtered out.
            
        Examples:
        ---------
        ```python
        # Internal usage (called by from_file and from_url)
        chunks = builder._split_text("This is a long text that needs to be split...")
        print(f"Created {len(chunks)} chunks")
        ```
        
        Notes:
        ------
        - Chunks are stripped of leading/trailing whitespace
        - Empty chunks are automatically filtered out
        - Word boundaries (spaces and newlines) are preferred break points
        - If no word boundary is found, text is broken at the exact size limit
        - The last chunk may be smaller than chunk_size
        - Overlap is applied between consecutive chunks
        """
        if len(text) <= self._chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Calculate end position for current chunk
            end = start + self._chunk_size
            
            # If this is not the last chunk, try to break at a word boundary
            if end < len(text):
                # Look for the last space or newline before the end
                last_space = text.rfind(' ', start, end)
                last_newline = text.rfind('\n', start, end)
                break_point = max(last_space, last_newline)
                
                # If we found a good break point, use it
                if break_point > start:
                    end = break_point + 1
            
            # Extract the chunk
            chunk = text[start:end].strip()
            if chunk:  # Only add non-empty chunks
                chunks.append(chunk)
            
            # Calculate next start position with overlap
            start = end - self._chunk_overlap
            if start >= len(text):
                break
        
        return chunks
        